"""
文书生成服务 - 按流程图实现
A: 用户进入文书生成 -> B: 是否已有已确认的案情分析?
  是: C 读取案情 -> D 匹配模板 -> E 展示 -> F 用户确认 -> G 填充 -> H 生成初稿 -> I 校验 -> K 生成Word -> L 下载
  否: M 提示提交案情 -> N 用户输入 -> O 标准化 -> P 是否足以支持? -> Q 补充问题 / S 推荐模板 -> ...
"""
import io
import json
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.db.connection import get_db
from app.services.llm_service import get_llm_answer


# 文书模板定义：template_key -> (title, 适用场景关键词)
DOC_TEMPLATES = [
    {"key": "divorce", "title": "离婚纠纷起诉状", "keywords": ["离婚", "婚姻", "财产分割", "子女抚养"]},
    {"key": "lending", "title": "民间借贷起诉状", "keywords": ["借贷", "借款", "欠款", "利息"]},
    {"key": "labor", "title": "劳动仲裁申请书", "keywords": ["劳动", "工资", "工伤", "辞退", "仲裁"]},
    {"key": "traffic", "title": "交通事故赔偿起诉状", "keywords": ["交通事故", "赔偿", "人身损害"]},
]


def get_sessions_with_confirmed_case(user_id: str, limit: int = 20) -> list:
    """获取用户已确认案情分析的会话列表"""
    db = get_db()
    cursor = db["sessions"].find(
        {"user_id": user_id, "case_status": "analysis_done"}
    ).sort("updated_at", -1).limit(limit)
    return [
        {
            "session_id": d["session_id"],
            "title": d.get("title", "案情咨询"),
            "case_summary": d.get("case_summary", ""),
            "analysis_result": d.get("analysis_result", {}),
            "updated_at": d.get("updated_at").isoformat() if d.get("updated_at") else None,
        }
        for d in cursor
    ]


def get_case_info_from_session(user_id: str, session_id: str) -> Optional[dict]:
    """从会话读取结构化案情信息与法律分析结果"""
    db = get_db()
    doc = db["sessions"].find_one({"session_id": session_id, "user_id": user_id})
    if not doc or doc.get("case_status") != "analysis_done":
        return None
    return {
        "case_input": doc.get("case_input", ""),
        "case_summary": doc.get("case_summary", ""),
        "analysis_result": doc.get("analysis_result", {}),
    }


async def match_templates_for_case(case_summary: str, analysis: Optional[dict] = None) -> list:
    """根据案情匹配推荐文书模板"""
    text = (case_summary or "") + " " + json.dumps(analysis or {}, ensure_ascii=False)
    text_lower = text.lower()
    scored = []
    for t in DOC_TEMPLATES:
        score = sum(1 for kw in t["keywords"] if kw in text_lower)
        scored.append((score, t))
    scored.sort(key=lambda x: -x[0])
    return [{"key": t["key"], "title": t["title"], "score": s} for s, t in scored if s > 0] or [
        {"key": t["key"], "title": t["title"], "score": 0} for t in DOC_TEMPLATES[:3]
    ]


def _parse_json_from_llm(text: str) -> Optional[dict]:
    text = (text or "").strip()
    for pattern in [r"```(?:json)?\s*([\s\S]*?)```", r"```\s*([\s\S]*?)```"]:
        m = re.search(pattern, text)
        if m:
            text = m.group(1).strip()
            break
    m = re.search(r"\{[\s\S]*\}", text)
    if m:
        text = m.group(0)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return None


async def check_case_sufficient_for_doc(case_input: str) -> dict:
    """检查案情是否足以支持文书生成，不足则返回补充问题"""
    prompt = f"""请判断以下案情描述是否足以支持生成法律文书（起诉状/申请书等）。

足以支持的标准：应包含当事人、主要事实、争议焦点、诉求等关键要素。
若不足，请生成 1-3 个补充问题。
若足够，请识别案件类型（如：离婚、借贷、劳动、交通事故等）并返回推荐模板。

请严格按 JSON 返回：
- 不足：{{"sufficient": false, "questions": ["问题1", "问题2"]}}
- 足够：{{"sufficient": true, "case_type": "离婚/借贷/劳动/交通事故", "templates": ["divorce"]}}

用户案情：
{case_input}
"""
    raw = await get_llm_answer(prompt, [])
    parsed = _parse_json_from_llm(raw)
    if parsed and parsed.get("sufficient"):
        templates = parsed.get("templates") or ["divorce"]
        return {"sufficient": True, "case_type": parsed.get("case_type", ""), "templates": templates}
    questions = []
    if parsed and isinstance(parsed.get("questions"), list):
        questions = parsed["questions"]
    if not questions:
        questions = ["请补充当事人信息", "请描述主要事实经过", "请说明您的诉求"]
    return {"sufficient": False, "questions": questions}


async def generate_document_draft(template_key: str, case_info: dict) -> tuple[str, list]:
    """
    生成文书初稿，返回 (draft_text, missing_fields)
    尽量从案情中提取具体信息填入文书，减少占位符
    """
    case_summary = case_info.get("case_summary") or case_info.get("case_input", "")
    analysis = case_info.get("analysis_result") or {}

    prompt = f"""请根据以下案情和法律分析，生成一份{_get_template_title(template_key)}的初稿。

案情摘要：
{case_summary}

法律分析：
{json.dumps(analysis, ensure_ascii=False)}

要求：
1. 格式规范，符合法律文书格式
2. 必须从案情中提取并填入具体信息：当事人姓名、日期（入职/离职/欠薪起止等）、金额、职务等，能确定的务必直接填入
3. 仅当案情中确实未提及的信息才使用占位符（如 ______ 或 ____年____月____日）
4. 包含：当事人信息、诉讼/仲裁请求、事实与理由等
5. 直接输出文书正文，不要输出 JSON 或说明文字
"""
    draft = await get_llm_answer(prompt, [])
    missing = _validate_draft(draft, template_key)
    return draft, missing


def _get_template_title(key: str) -> str:
    for t in DOC_TEMPLATES:
        if t["key"] == key:
            return t["title"]
    return "法律文书"


def _validate_draft(draft: str, template_key: str) -> list:
    """校验文书初稿，返回缺失的关键信息"""
    missing = []
    placeholders = ["________", "______", "（待填写）", "（  ）"]
    if any(p in draft for p in placeholders):
        missing.append("部分当事人信息待填写")
    required = ["原告", "被告", "诉讼请求", "事实与理由"]
    for r in required:
        if r not in draft and "申请" not in draft:
            if "原告" in required and template_key == "labor":
                continue
            missing.append(f"建议补充「{r}」部分")
    return missing[:5]  # 最多返回5项


def create_word_document(draft: str, template_key: str, filename: str) -> bytes:
    """将文书初稿导出为 Word 文件"""
    doc = Document()
    title = _get_template_title(template_key)
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    # 正文
    for line in draft.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def save_document_for_download(user_id: str, template_key: str, draft: str) -> str:
    """保存文书到临时存储，返回文档 ID"""
    db = get_db()
    doc_id = str(uuid.uuid4())
    content = create_word_document(draft, template_key, f"{template_key}.docx")
    db["document_downloads"].insert_one({
        "doc_id": doc_id,
        "user_id": user_id,
        "template_key": template_key,
        "content": content,
        "created_at": datetime.utcnow(),
    })
    return doc_id


def get_document_for_download(user_id: str, doc_id: str) -> Optional[bytes]:
    """获取待下载的文书内容"""
    db = get_db()
    doc = db["document_downloads"].find_one({"doc_id": doc_id, "user_id": user_id})
    if not doc:
        return None
    return doc.get("content")


_TEMPLATE_TITLES = {
    "divorce": "离婚纠纷起诉状",
    "lending": "民间借贷起诉状",
    "labor": "劳动仲裁申请书",
    "traffic": "交通事故赔偿起诉状",
}


def delete_document(user_id: str, doc_id: str) -> bool:
    """删除文书记录，校验归属"""
    db = get_db()
    result = db["document_downloads"].delete_one({"doc_id": doc_id, "user_id": user_id})
    return result.deleted_count > 0


def get_user_documents(user_id: str, limit: int = 50) -> list:
    """获取用户的文书生成记录列表"""
    db = get_db()
    cursor = db["document_downloads"].find({"user_id": user_id}).sort("created_at", -1).limit(limit)
    return [
        {
            "doc_id": d["doc_id"],
            "template_key": d.get("template_key", ""),
            "title": _TEMPLATE_TITLES.get(d.get("template_key", ""), "文书"),
            "created_at": d.get("created_at").isoformat() if d.get("created_at") else None,
        }
        for d in cursor
    ]
